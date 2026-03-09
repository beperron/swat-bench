#!/usr/bin/env python3
"""Generate a human-readable .docx of the SWAT-Bench exam (55 tasks)."""

import json
import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.join(os.path.dirname(__file__), "..")
TASKS_DIR = os.path.join(BASE, "tasks")
OUTPUT = os.path.join(BASE, "SWAT-Bench_Exam.docx")

DOMAIN_ORDER = [
    "Data Cleaning & Validation",
    "Data Preparation & Transformation",
    "Descriptive Statistics & Measurement",
    "Inferential Statistics",
    "Applied Social Work Analytics",
    "Text & Natural Language Processing",
]

DIFFICULTY_ORDER = {"Foundational": 0, "Intermediate": 1, "Advanced": 2}

CHECK_TYPE_LABELS = {
    "execution": "Code executes without error",
    "exact": "Exact match",
    "numeric": "Numeric match (within tolerance)",
    "range": "Value within range",
    "regex": "Pattern match (regex)",
    "file_exists": "Output file exists",
}


def load_tasks():
    """Load all tasks grouped by domain."""
    domains = {}
    for name in sorted(os.listdir(TASKS_DIR)):
        task_dir = os.path.join(TASKS_DIR, name)
        checks_path = os.path.join(task_dir, "expected", "checks.json")
        prompt_path = os.path.join(task_dir, "prompt.md")
        if not os.path.isfile(checks_path) or not os.path.isfile(prompt_path):
            continue
        with open(checks_path) as f:
            checks = json.load(f)
        with open(prompt_path) as f:
            prompt = f.read()
        domain = checks.get("domain", "Unknown")
        if domain not in domains:
            domains[domain] = []
        domains[domain].append({
            "dir_name": name,
            "test_id": checks.get("test_id", name),
            "task_name": checks.get("task_name", name),
            "domain": domain,
            "difficulty": checks.get("difficulty", "Unknown"),
            "total_points": checks.get("total_points", 0),
            "checks": checks.get("checks", []),
            "prompt": prompt,
        })
    return domains


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shading_elem)


def format_check_description(check):
    """Create a human-readable description of a scoring check."""
    ctype = check["type"]
    label = check.get("label", "")
    points = check.get("points", 0)

    if ctype == "execution":
        return f"Code runs successfully ({points} pt)"

    if ctype == "file_exists":
        fname = check.get("filename", "?")
        return f'Output file "{fname}" exists ({points} pt)'

    # Build the description from pattern + expected
    pattern = check.get("pattern", "")
    expected = check.get("expected", "")
    tolerance = check.get("tolerance", None)
    expected_pattern = check.get("expected_pattern", "")
    min_val = check.get("min", "")
    max_val = check.get("max", "")

    # Clean up regex pattern for display
    display_pattern = pattern
    # Replace capture groups with placeholder
    display_pattern = re.sub(r'\(\\d\+\)', '<N>', display_pattern)
    display_pattern = re.sub(r'\(-?\\d\+\\.\\d\+\)', '<X.X>', display_pattern)
    display_pattern = re.sub(r'\(-?\\d\+\\.?\\d*\)', '<X>', display_pattern)
    display_pattern = re.sub(r'\([\.\*\+\?\[\]\\dawsS]+?\)', '<value>', display_pattern)
    # Remove remaining escape chars for display
    display_pattern = display_pattern.replace("\\", "")

    if ctype == "exact":
        return f'{display_pattern} = {expected} ({points} pt)'
    elif ctype == "numeric":
        tol_str = f" (tolerance: {tolerance})" if tolerance else ""
        return f'{display_pattern} = {expected}{tol_str} ({points} pt)'
    elif ctype == "range":
        return f'{display_pattern} in [{min_val}, {max_val}] ({points} pt)'
    elif ctype == "regex":
        return f'Output matches: {expected_pattern or pattern} ({points} pt)'
    else:
        return f'{label}: {ctype} ({points} pt)'


def add_prompt_content(doc, prompt_text):
    """Add the prompt.md content with basic markdown rendering."""
    lines = prompt_text.strip().split("\n")
    in_code_block = False
    code_lines = []

    for line in lines:
        # Code block toggle
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block — render accumulated code
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.style = doc.styles['No Spacing']
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
                pf = p.paragraph_format
                pf.space_before = Pt(2)
                pf.space_after = Pt(2)
                pf.left_indent = Cm(1)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            continue

        # Headings
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            continue
        if stripped.startswith("# "):
            # Skip top-level heading (we use our own)
            continue

        # Bullet points
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            # Handle inline bold
            parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(10)
                elif part.startswith("`") and part.endswith("`"):
                    run = p.add_run(part[1:-1])
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(10)
                elif part.startswith("`") and part.endswith("`"):
                    run = p.add_run(part[1:-1])
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            continue

        # Regular paragraph with inline formatting
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*|`.*?`)', stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(10)
            elif part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            else:
                run = p.add_run(part)
                run.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title page ──
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SWAT-Bench")
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Structured Data Analysis Benchmark\nfor Local LLMs in Social Work Research")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary.add_run("55 Tasks  |  435 Points  |  7 Domains")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("Complete Exam Reference — Prompts & Scoring Criteria")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ── Table of Contents ──
    toc_heading = doc.add_heading("Table of Contents", level=1)

    domains = load_tasks()
    task_num = 0
    for domain_name in DOMAIN_ORDER:
        tasks = domains.get(domain_name, [])
        if not tasks:
            continue
        tasks.sort(key=lambda t: t["dir_name"])
        domain_points = sum(t["total_points"] for t in tasks)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{domain_name} ({len(tasks)} tasks, {domain_points} points)")
        run.bold = True
        run.font.size = Pt(11)

        for task in tasks:
            task_num += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            diff = task["difficulty"]
            run = p.add_run(
                f'{task["test_id"]}  {task["task_name"]}  '
                f'[{diff}, {task["total_points"]} pts]'
            )
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_page_break()

    # ── Domain sections ──
    task_num = 0
    for domain_idx, domain_name in enumerate(DOMAIN_ORDER, 1):
        tasks = domains.get(domain_name, [])
        if not tasks:
            continue
        tasks.sort(key=lambda t: t["dir_name"])
        domain_points = sum(t["total_points"] for t in tasks)

        # Domain heading
        doc.add_heading(f"Domain {domain_idx}: {domain_name}", level=1)
        p = doc.add_paragraph()
        run = p.add_run(f"{len(tasks)} tasks  |  {domain_points} points")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.italic = True

        # Horizontal rule
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)

        for task in tasks:
            task_num += 1

            # ── Task header ──
            doc.add_heading(
                f'{task["test_id"]}: {task["task_name"]}',
                level=2,
            )

            # Metadata line
            diff_colors = {
                "Foundational": "2E7D32",
                "Intermediate": "E65100",
                "Advanced": "B71C1C",
            }
            p = doc.add_paragraph()
            run = p.add_run(f'Difficulty: ')
            run.font.size = Pt(10)
            run.bold = True
            run = p.add_run(f'{task["difficulty"]}')
            run.font.size = Pt(10)
            color = diff_colors.get(task["difficulty"], "333333")
            run.font.color.rgb = RGBColor(
                int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
            )
            run.bold = True
            run = p.add_run(f'    |    Points: ')
            run.font.size = Pt(10)
            run.bold = True
            run = p.add_run(f'{task["total_points"]}')
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

            # ── Prompt section ──
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            run = p.add_run("Prompt")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

            # Thin separator
            p = doc.add_paragraph()
            pBdr = p._element.get_or_add_pPr()
            border = pBdr.makeelement(qn("w:pBdr"), {})
            bottom = border.makeelement(qn("w:bottom"), {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "1",
                qn("w:color"): "CCCCCC",
            })
            border.append(bottom)
            pBdr.append(border)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)

            add_prompt_content(doc, task["prompt"])

            # ── Scoring criteria ──
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            run = p.add_run("Scoring Criteria")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

            # Scoring table
            checks = task["checks"]
            table = doc.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.style = 'Table Grid'

            # Header row
            hdr = table.rows[0]
            for i, text in enumerate(["#", "Criterion", "Points"]):
                cell = hdr.cells[i]
                cell.text = ""
                run = cell.paragraphs[0].add_run(text)
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
                set_cell_shading(cell, "1A3C6E")

            for idx, check in enumerate(checks, 1):
                row = table.add_row()
                # Number
                row.cells[0].text = ""
                run = row.cells[0].paragraphs[0].add_run(str(idx))
                run.font.size = Pt(9)
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Description
                desc_text = format_check_description(check)
                row.cells[1].text = ""
                run = row.cells[1].paragraphs[0].add_run(desc_text)
                run.font.size = Pt(9)
                # Points
                pts = check.get("points", 0)
                row.cells[2].text = ""
                run = row.cells[2].paragraphs[0].add_run(str(pts))
                run.font.size = Pt(9)
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Alternate row shading
                if idx % 2 == 0:
                    for cell in row.cells:
                        set_cell_shading(cell, "F2F6FA")

            # Total row
            total_row = table.add_row()
            total_row.cells[0].text = ""
            total_row.cells[1].text = ""
            run = total_row.cells[1].paragraphs[0].add_run("Total")
            run.bold = True
            run.font.size = Pt(9)
            total_row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            total_row.cells[2].text = ""
            run = total_row.cells[2].paragraphs[0].add_run(str(task["total_points"]))
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
            total_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for cell in total_row.cells:
                set_cell_shading(cell, "E8EDF5")

            # Set column widths
            for row in table.rows:
                row.cells[0].width = Cm(1.2)
                row.cells[1].width = Cm(12)
                row.cells[2].width = Cm(2)

            # Page break after each task (except last in domain)
            if task != tasks[-1] or domain_name != DOMAIN_ORDER[-1]:
                doc.add_page_break()

    # Save
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Tasks: {task_num}")


if __name__ == "__main__":
    build_document()
